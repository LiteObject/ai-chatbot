"""
Document handling functionality for the AI Chatbot.
Handles file uploads, processing, and indexing using LlamaIndex.
"""
import os
import streamlit as st
from typing import Optional, Dict, Any
from pathlib import Path

# Runtime imports with fallbacks
try:
    from llama_index.core import VectorStoreIndex, Document, StorageContext, Settings  # type: ignore
    from llama_index.llms.openai import OpenAI  # type: ignore
    from llama_index.embeddings.openai import OpenAIEmbedding  # type: ignore
    from llama_index.vector_stores.chroma import ChromaVectorStore  # type: ignore
    import chromadb  # type: ignore
    import PyPDF2  # type: ignore
    import docx  # type: ignore
    IMPORTS_AVAILABLE = True
except ImportError as e:
    # Try alternative import for older versions
    try:
        from llama_index import VectorStoreIndex, Document, StorageContext  # type: ignore
        try:
            from llama_index import ServiceContext  # type: ignore
        except ImportError:
            ServiceContext = None
        from llama_index.llms import OpenAI  # type: ignore
        from llama_index.embeddings import OpenAIEmbedding  # type: ignore
        from llama_index.vector_stores import ChromaVectorStore  # type: ignore
        import chromadb  # type: ignore
        import PyPDF2  # type: ignore
        import docx  # type: ignore
        Settings = None  # Use ServiceContext for older versions
        IMPORTS_AVAILABLE = True
    except ImportError as e2:
        # Runtime imports not available
        IMPORTS_AVAILABLE = False
        import_error = str(e2)

from config.settings import settings
from src.utils import (
    get_file_hash,
    format_file_size,
    is_valid_file_type,
    sanitize_filename,
    get_timestamp,
    ensure_directory_exists
)


class DocumentHandler:
    """Handles document processing and indexing."""

    def __init__(self):
        """Initialize the document handler."""
        self.upload_dir = settings.UPLOAD_DIR
        self.data_dir = settings.DATA_DIR
        ensure_directory_exists(self.upload_dir)
        ensure_directory_exists(self.data_dir)

        # Initialize ChromaDB
        self.chroma_client = None
        self.vector_store = None
        self.storage_context = None
        self.document_index = None

        if IMPORTS_AVAILABLE:
            self._initialize_vector_store()
            self._initialize_service_context()
        else:
            # Store for later display when features are used
            self.import_error = globals().get(
                'import_error', 'LlamaIndex libraries not available')

    def _initialize_vector_store(self):
        """Initialize ChromaDB vector store."""
        if not IMPORTS_AVAILABLE:
            st.error(
                "Cannot initialize vector store: Required libraries not available")
            return

        try:
            # Connect to ChromaDB running in Docker
            self.chroma_client = chromadb.HttpClient(  # type: ignore
                host="localhost",
                port=8000
            )

            chroma_collection = self.chroma_client.get_or_create_collection(
                "documents")
            self.vector_store = ChromaVectorStore(  # type: ignore
                chroma_collection=chroma_collection)
            self.storage_context = StorageContext.from_defaults(  # type: ignore
                vector_store=self.vector_store)
        except (ConnectionError, OSError, ImportError, AttributeError) as e:
            st.error(f"Failed to initialize vector store: {e}")
            st.error("Make sure ChromaDB is running: docker-compose up -d chromadb")

    def _initialize_service_context(self):
        """Initialize LlamaIndex settings."""
        if not IMPORTS_AVAILABLE:
            st.error(
                "Cannot initialize service context: Required libraries not available")
            return

        try:
            temperature = st.session_state.get(
                'temperature', settings.DEFAULT_TEMPERATURE)
            model = st.session_state.get('model', settings.DEFAULT_MODEL)

            # Configure global settings (only if Settings is available)
            if Settings is not None:
                Settings.llm = OpenAI(  # type: ignore
                    model=model,
                    temperature=temperature,
                    api_key=settings.OPENAI_API_KEY
                )
                Settings.embed_model = OpenAIEmbedding(  # type: ignore
                    api_key=settings.OPENAI_API_KEY)
                Settings.chunk_size = settings.DEFAULT_CHUNK_SIZE
                Settings.chunk_overlap = settings.DEFAULT_CHUNK_OVERLAP

        except (ImportError, AttributeError, ValueError) as e:
            st.error(f"Failed to initialize service context: {e}")

    def validate_file(self, uploaded_file) -> Dict[str, Any]:
        """Validate uploaded file."""
        validation_result = {
            "is_valid": False,
            "error_message": "",
            "file_info": {}
        }

        # Check file type
        if not is_valid_file_type(uploaded_file.name, settings.SUPPORTED_FILE_TYPES):
            validation_result[
                "error_message"] = f"Unsupported file type. Supported types: {', '.join(settings.SUPPORTED_FILE_TYPES)}"
            return validation_result

        # Check file size
        file_size_mb = uploaded_file.size / (1024 * 1024)
        if file_size_mb > settings.MAX_FILE_SIZE_MB:
            validation_result["error_message"] = f"File size ({file_size_mb:.1f}MB) exceeds limit of {settings.MAX_FILE_SIZE_MB}MB"
            return validation_result

        # Check total documents limit
        if len(self.get_uploaded_files()) >= settings.MAX_DOCUMENTS:
            validation_result[
                "error_message"] = f"Maximum number of documents ({settings.MAX_DOCUMENTS}) reached"
            return validation_result

        validation_result["is_valid"] = True
        validation_result["file_info"] = {
            "name": uploaded_file.name,
            "size": uploaded_file.size,
            "size_formatted": format_file_size(uploaded_file.size),
            "type": uploaded_file.type
        }

        return validation_result

    def save_uploaded_file(self, uploaded_file) -> Optional[str]:
        """Save uploaded file to uploads directory."""
        try:
            file_content = uploaded_file.read()
            file_hash = get_file_hash(file_content)
            sanitized_name = sanitize_filename(uploaded_file.name)
            file_path = os.path.join(
                self.upload_dir, f"{file_hash}_{sanitized_name}")

            with open(file_path, "wb") as f:
                f.write(file_content)

            # Store file metadata
            file_metadata = {
                "original_name": uploaded_file.name,
                "file_path": file_path,
                "file_hash": file_hash,
                "size": uploaded_file.size,
                "upload_time": get_timestamp(),
                "processed": False
            }

            # Add to session state
            if "uploaded_files" not in st.session_state:
                st.session_state.uploaded_files = {}

            st.session_state.uploaded_files[file_hash] = file_metadata

            return file_path

        except (OSError, IOError, PermissionError) as e:
            st.error(f"Failed to save file: {e}")
            return None

    def process_document(self, file_path: str, file_hash: str) -> bool:
        """Process document and add to index."""
        if not IMPORTS_AVAILABLE:
            st.error("Cannot process document: Required libraries not available")
            return False

        try:
            # Extract text based on file type
            file_extension = Path(file_path).suffix.lower()

            if file_extension == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                text = self._extract_docx_text(file_path)
            elif file_extension == '.txt':
                text = self._extract_txt_text(file_path)
            else:
                st.error(f"Unsupported file type: {file_extension}")
                return False

            if not text.strip():
                st.error("No text content found in the document")
                return False

            # Create document object
            document = Document(  # type: ignore
                text=text,
                metadata={
                    "file_path": file_path,
                    "file_hash": file_hash,
                    "file_name": st.session_state.uploaded_files[file_hash]["original_name"],
                    "upload_time": st.session_state.uploaded_files[file_hash]["upload_time"]
                }
            )

            # Add to index
            if self.document_index is None:
                self.document_index = VectorStoreIndex.from_documents(  # type: ignore
                    [document],
                    storage_context=self.storage_context
                )
            else:
                self.document_index.insert(document)

            # Update metadata
            st.session_state.uploaded_files[file_hash]["processed"] = True

            return True

        except (ImportError, AttributeError, ValueError, KeyError) as e:
            st.error(f"Failed to process document: {e}")
            return False

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        if not IMPORTS_AVAILABLE:
            st.error("PDF reading not available: PyPDF2 library not installed")
            return ""

        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)  # type: ignore
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except (OSError, IOError, AttributeError) as e:
            st.error(f"Failed to extract PDF text: {e}")
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        if not IMPORTS_AVAILABLE:
            st.error(
                "DOCX reading not available: python-docx library not installed")
            return ""

        text = ""
        try:
            doc = docx.Document(file_path)  # type: ignore
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except (OSError, IOError, AttributeError) as e:
            st.error(f"Failed to extract DOCX text: {e}")
        return text

    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            except (OSError, IOError, UnicodeDecodeError) as e:
                st.error(f"Failed to extract TXT text: {e}")
        except (OSError, IOError) as e:
            st.error(f"Failed to extract TXT text: {e}")
        return text

    def delete_document(self, file_hash: str) -> bool:
        """Delete document from storage and index."""
        try:
            if file_hash not in st.session_state.uploaded_files:
                return False

            file_metadata = st.session_state.uploaded_files[file_hash]
            file_path = file_metadata["file_path"]

            # Remove file from disk
            if os.path.exists(file_path):
                os.remove(file_path)

            # Remove from session state
            del st.session_state.uploaded_files[file_hash]

            # Note: ChromaDB doesn't easily support removing specific documents
            # For production, you might want to rebuild the index or use a different approach

            return True

        except (OSError, KeyError) as e:
            st.error(f"Failed to delete document: {e}")
            return False

    def get_uploaded_files(self) -> Dict[str, Dict]:
        """Get list of uploaded files."""
        return st.session_state.get("uploaded_files", {})

    def clear_all_documents(self) -> bool:
        """Clear all documents and reset index."""
        try:
            # Remove all files
            for file_hash in list(st.session_state.get("uploaded_files", {}).keys()):
                self.delete_document(file_hash)

            # Reset index
            self.document_index = None

            # Clear uploads directory
            for filename in os.listdir(self.upload_dir):
                file_path = os.path.join(self.upload_dir, filename)
                if os.path.isfile(file_path):
                    os.remove(file_path)

            return True

        except (OSError, IOError, PermissionError) as e:
            st.error(f"Failed to clear documents: {e}")
            return False

    def get_query_engine(self):
        """Get query engine for document search."""
        if not IMPORTS_AVAILABLE:
            st.error("Cannot create query engine: Required libraries not available")
            return None

        if self.document_index is None:
            return None
        return self.document_index.as_query_engine()
